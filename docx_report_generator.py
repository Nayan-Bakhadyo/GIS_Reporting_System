from tkinter import simpledialog
import tkinter as tk
from docx import Document
from docx.shared import Inches
import os
from PIL import Image
import time

image_path = '/Users/Nayan/Desktop/work/GIS_Reporting_System/Temp/output/images/test.png'

output_pdf = '/Users/Nayan/Desktop/work/GIS_Reporting_System/Temp/output/pdfs/document.pdf'

from PIL import Image

def generate_docx():
    ward = simpledialog.askinteger("Ward Number", "Enter the Ward number:")
    tole = simpledialog.askinteger("Tole Number", "Enter Tole number:")

    output_docx = ('/Users/Nayan/Desktop/work/GIS_Reporting_System/Temp/output/docx/Ward_'+str(ward)+'_Tole_'+str(tole)+'.docx')

    image_path = os.path.join(os.environ["GIS_REPORTING_SYSTEM"],
                              "GIS_Reporting_System/Temp/output/images/")
    images = [file for file in os.listdir(image_path) if file.endswith(".png")]

    print('')
    print(f'Generating Report for Ward {ward} - Tole {tole}!!!!!!')
    # Create a new Word document
    doc = Document()
    # Add a heading
    doc.add_heading('Report for Ward '+ str(ward) + ' - Tole ' + str(tole) , level=1)

    # Create a table with two columns and four rows
    table = doc.add_table(rows=6, cols=2)
    table.autofit = True  # Disable auto resizing of the table cells

    # Set the width of the first column to half of the page width
    table.columns[0].width = Inches(4)
    images = ['shapefile_map.png', 'Building_points.png', 'Base_Transceiver_Station_BTS_Tower.png', 'Sewer_Network.png','Existing_KV-WaterSupply_Network.png','KUKL_Proposed_Pipeline_Network.png']

    # Add images to the first row
    for i in range(2):
        row = table.rows[0]
        cell = row.cells[i]
        cell_paragraph = cell.paragraphs[0]
        run = cell_paragraph.add_run()
        image = Image.open(os.path.join(image_path, images[i]))
        width, height = image.size
        aspect_ratio = width / height
        # Adjust the width of the cell based on the aspect ratio of the image
        adjusted_width = cell.width - Inches(0.2)  # Assuming a margin of 0.1 inch on each side
        # Calculate the number of lines needed to fit the image
        #num_lines = int(adjusted_width / aspect_ratio / Inches(1.7))  # Assuming average font size of 12pt
        num_lines = 0
        # Set the height of the cell to accommodate the image
        cell_paragraph.add_run("\n" * num_lines)  # Add empty lines to adjust cell height
        run.add_picture(os.path.join(image_path, images[i]), width=adjusted_width, height=adjusted_width / aspect_ratio)

    # Add text to the first row first column
    cell = table.cell(1, 0)
    cell.text = f"Boundary Area of Ward {ward} Tole {tole}"
    # Adjust the height of the row to fit the text
    cell.height = Inches(0.1)  # You may need to adjust the height based on your text content
    cell_paragraph.add_run("\n" * num_lines)  # Add empty lines to adjust cell height

    # Add text to the first row second column
    cell = table.cell(1, 1)
    cell.text = f"Building Points of Ward {ward} Tole {tole}"
    # Adjust the height of the row to fit the text
    cell.height = Inches(0.1)  # You may need to adjust the height based on your text content
    cell_paragraph.add_run("\n" * num_lines)  # Add empty lines to adjust cell height

    # Add images to the third row
    for i in range(2, 4):
        row = table.rows[2]
        cell = row.cells[i-2]
        cell_paragraph = cell.paragraphs[0]
        run = cell_paragraph.add_run()
        image = Image.open(os.path.join(image_path, images[i]))
        width, height = image.size
        aspect_ratio = width / height
        # Adjust the width of the cell based on the aspect ratio of the image
        adjusted_width = cell.width - Inches(0.2)  # Assuming a margin of 0.1 inch on each side
        # Calculate the number of lines needed to fit the image
        #num_lines = int(adjusted_width / aspect_ratio / Inches(0.2))  # Assuming average font size of 12pt
        num_lines = 0
        # Set the height of the cell to accommodate the image
        cell_paragraph.add_run("\n" * num_lines)  # Add empty lines to adjust cell height
        run.add_picture(os.path.join(image_path, images[i]), width=adjusted_width, height=adjusted_width / aspect_ratio)

    # Add text to the second row first column
    cell = table.cell(3, 0)
    cell.text = "GSM Tower"
    # Adjust the height of the row to fit the text
    cell.height = Inches(0.1)  # You may need to adjust the height based on your text content

    # Add text to the second row first column
    cell = table.cell(3, 1)
    cell.text = "Sewer Network"
    # Adjust the height of the row to fit the text
    cell.height = Inches(0.1)  # You may need to adjust the height based on your text content

    # Add images to the third row
    for i in range(4, 6):
        row = table.rows[4]
        cell = row.cells[i - 4]
        cell_paragraph = cell.paragraphs[0]
        run = cell_paragraph.add_run()
        image = Image.open(os.path.join(image_path, images[i]))
        width, height = image.size
        aspect_ratio = width / height
        # Adjust the width of the cell based on the aspect ratio of the image
        adjusted_width = cell.width - Inches(0.2)  # Assuming a margin of 0.1 inch on each side
        # Calculate the number of lines needed to fit the image
        # num_lines = int(adjusted_width / aspect_ratio / Inches(0.2))  # Assuming average font size of 12pt
        num_lines = 0
        # Set the height of the cell to accommodate the image
        cell_paragraph.add_run("\n" * num_lines)  # Add empty lines to adjust cell height
        run.add_picture(os.path.join(image_path, images[i]), width=adjusted_width, height=adjusted_width / aspect_ratio)

    # Add text to the second row first column
    cell = table.cell(5, 0)
    cell.text = "Existing KV water supply network"
    # Adjust the height of the row to fit the text
    cell.height = Inches(0.1)  # You may need to adjust the height based on your text content

    # Add text to the second row first column
    cell = table.cell(5, 1)
    cell.text = "KUKL Proposed pipeline network"
    # Adjust the height of the row to fit the text
    cell.height = Inches(0.1)  # You may need to adjust the height based on your text content


    # Save the document
    doc.save(output_docx)


#generate_docx()